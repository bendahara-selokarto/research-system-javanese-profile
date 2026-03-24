from __future__ import annotations

import argparse
from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document

from research_system.utils import (
    EVENT_GUIDELINES,
    SELAPAN_CYCLE_DAYS,
    compatibility_result,
    get_watak_profile,
    hari_baik_advice,
    javanese_day_profile,
)


DEFAULT_EVENTS = list(EVENT_GUIDELINES.keys())


def _parse_date(raw: str) -> date:
    return date.fromisoformat(raw)


def _artifact_stem(target: date, partner: date | None = None) -> str:
    stem = f"{target:%Y-%m-%d}"
    if partner is not None:
        stem = f"{stem}-partner-{partner:%Y-%m-%d}"
    return stem


def _next_artifact_path(output_dir: Path, stem: str) -> Path:
    candidate = output_dir / f"{stem}.docx"
    if not candidate.exists():
        return candidate

    suffix = 2
    while True:
        candidate = output_dir / f"{stem}-{suffix}.docx"
        if not candidate.exists():
            return candidate
        suffix += 1


def _human_join(values: Iterable[str]) -> str:
    parts = tuple(values)
    if len(parts) == 1:
        return parts[0]
    if len(parts) == 2:
        return f"{parts[0]} dan {parts[1]}"
    return ", ".join(parts[:-1]) + f", dan {parts[-1]}"


def write_profile_docx(
    target_date: date | str,
    partner_date: date | str | None = None,
    events: Iterable[str] | None = None,
    output_dir: Path | None = None,
) -> Path:
    target = _parse_date(target_date) if isinstance(target_date, str) else target_date
    partner = _parse_date(partner_date) if isinstance(partner_date, str) else partner_date
    profile = javanese_day_profile(target)
    output_base = (output_dir or Path("output")).expanduser()
    output_base.mkdir(parents=True, exist_ok=True)
    artifact = _next_artifact_path(output_base, _artifact_stem(target, partner))

    document = Document()
    document.add_heading(f"Profil Hari Jawa - {target:%Y-%m-%d}", level=1)
    document.add_paragraph(profile.summary)

    document.add_heading("Identitas dasar", level=2)
    table = document.add_table(rows=0, cols=2)
    identity_rows = [
        ("Hari", profile.identity.hari),
        ("Dinapitu", profile.identity.dinapitu),
        ("Pasaran", profile.identity.pasaran),
    ]
    if profile.identity.javanese_date is not None and profile.identity.hijri_date is not None:
        identity_rows.extend(
            [
                ("Tanggal Jawa", profile.identity.javanese_date.formatted),
                ("Tanggal Hijriyah", profile.identity.hijri_date.formatted),
            ]
        )
    identity_rows.extend(
        [
            ("Weton Jawa", profile.identity.weton_jawa),
            ("Wuku", profile.identity.wuku),
            ("Neptu total / jenjem", str(profile.identity.neptu_total)),
            ("Watak singkat", get_watak_profile(target)),
        ]
    )
    for label, value in identity_rows:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value

    if profile.identity.javanese_date is not None and profile.identity.hijri_date is not None:
        document.add_paragraph(
            "Tanggal Jawa pada sistem ini mengikuti hitungan lunar yang disejajarkan dengan kalender Hijriyah."
        )

    document.add_heading("Siklus tahun Jawa", level=2)
    if profile.identity.year_cycle is not None:
        year_cycle = profile.identity.year_cycle
        year_table = document.add_table(rows=0, cols=2)
        for label, value in (
            ("Tahun Jawa", str(year_cycle.year_number)),
            ("Nomor lanjutan Saka", str(year_cycle.saka_continuity_year)),
            ("Nama taun", year_cycle.year_name),
            ("Windu", year_cycle.windu_name),
            ("Tahun ke dalam windu", str(year_cycle.windu_year_number)),
            ("Jenis taun", f"{year_cycle.year_type} ({year_cycle.year_length_days} hari)"),
            (
                "Kurup",
                f"{year_cycle.kurup_code} ({year_cycle.kurup_name}, {year_cycle.kurup_start_year}-{year_cycle.kurup_end_year} AJ)",
            ),
        ):
            row = year_table.add_row().cells
            row[0].text = label
            row[1].text = value

        document.add_paragraph(
            "Nomor tahun Jawa Sultan Agungan meneruskan angka tahun Saka, "
            "tetapi perhitungan tahunnya mengikuti pola lunar yang diselaraskan dengan Hijriah."
        )
    else:
        document.add_paragraph(
            "Hitungan exact nama taun, windu, dan kurup belum ditampilkan untuk tanggal ini karena berada di luar rentang kurup baku yang diimplementasikan sistem."
        )

    naga_dina = profile.identity.naga_dina
    document.add_heading("Naga dina", level=2)
    document.add_paragraph(naga_dina.summary)
    naga_table = document.add_table(rows=0, cols=2)
    for label, value in (
        ("Varian default sistem", naga_dina.default_variant_label),
        ("Posisi hari", f"{profile.identity.dinapitu}: {_human_join(naga_dina.day_directions)}"),
        ("Posisi pasaran", f"{profile.identity.pasaran}: {naga_dina.pasaran_direction}"),
        ("Varian pembanding", "Boyongan neptu berputar"),
        ("Arah dari jumlah neptu", f"{naga_dina.neptu_cycle_total} -> {naga_dina.neptu_cycle_direction}"),
    ):
        row = naga_table.add_row().cells
        row[0].text = label
        row[1].text = value

    for variant in naga_dina.variants:
        document.add_paragraph(f"{variant.label}: {variant.basis} {variant.note} [{variant.source_label}]")

    document.add_heading("Selapan dan wetonan berikutnya", level=2)
    document.add_paragraph(
        f"Selapan ke-{profile.selapan_day} dari {SELAPAN_CYCLE_DAYS}-hari; "
        f"weton berikutnya: {profile.next_weton_date.isoformat()}."
    )
    for idx, next_date in enumerate(profile.next_three_weton_dates, start=1):
        document.add_paragraph(f"#{idx}: {next_date.isoformat()}")

    document.add_heading("Kegunaan budaya", level=2)
    for cultural_use in profile.common_uses:
        paragraph = document.add_paragraph()
        paragraph.add_run(cultural_use.category.replace("_", " ").capitalize() + ": ").bold = True
        paragraph.add_run(f"{cultural_use.description} ({cultural_use.example_question})")
        if cultural_use.requires_additional_input:
            paragraph.add_run(" [Butuh input tambahan]").italic = True

    document.add_heading("Kecocokan jodoh", level=2)
    if partner:
        compatibility = compatibility_result(target, partner)
        document.add_paragraph(
            f"{compatibility.first_weton} + {compatibility.second_weton} = "
            f"{compatibility.jenjem_sum} ({compatibility.category}). {compatibility.recommendation}"
        )
    else:
        document.add_paragraph("Tambahkan --partner-date untuk melihat jenjem pasangan.")

    document.add_heading("Hari baik", level=2)
    events_to_check = events or DEFAULT_EVENTS
    for event in events_to_check:
        advice = hari_baik_advice(target, event)
        document.add_paragraph(
            f"{event.title()}: {'Baik' if advice.is_good else 'Tidak baik'} - {advice.reason} ({advice.note})"
        )

    document.save(artifact)
    return artifact


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate a Word document that profiles a Javanese hari or weton."
    )
    parser.add_argument(
        "--date",
        default=date.today().isoformat(),
        type=_parse_date,
        help="Tanggal Masehi dalam format YYYY-MM-DD.",
    )
    parser.add_argument(
        "--partner-date",
        type=_parse_date,
        help="Tanggal Masehi pasangan untuk menghitung jenjem dua orang.",
    )
    parser.add_argument(
        "--events",
        nargs="*",
        choices=DEFAULT_EVENTS,
        help="Jenis acara yang ingin diperiksa hari baiknya.",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=Path("output"),
        help="Direktori tempat menyimpan file .docx.",
    )

    args = parser.parse_args()
    artifact = write_profile_docx(
        args.date,
        partner_date=args.partner_date,
        events=args.events,
        output_dir=args.output_dir,
    )
    print(f"{artifact} telah dibuat.")


if __name__ == "__main__":
    main()
