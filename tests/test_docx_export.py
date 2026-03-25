from __future__ import annotations

from contextlib import contextmanager
from datetime import date
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Iterator

from docx import Document

from research_system.commands.javanese_profile import write_profile_docx


@contextmanager
def output_dir() -> Iterator[Path]:
    base_dir = Path("output")
    base_dir.mkdir(exist_ok=True)
    with TemporaryDirectory(dir=base_dir) as tmp_dir:
        yield Path(tmp_dir)



def _document_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    parts.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    return "\n".join(parts)



def test_docx_export_contains_summary() -> None:
    with output_dir() as artifact_dir:
        artifact = write_profile_docx(
            target_date=date(1990, 4, 25),
            partner_date=date(2025, 1, 14),
            events=("nikah",),
            output_dir=artifact_dir,
        )

        doc = Document(artifact)
        text = _document_text(doc)
        assert "Rebo Pon" in text
        assert "nikah" in text.lower()
        assert "selamatan" in text.lower()
        assert "jenjem" in text.lower()
        assert "kurup" in text.lower()
        assert "saka" in text.lower()
        assert "pranata mangsa" in text.lower()
        assert "dhesta" in text.lower()
        assert "naga dina" in text.lower()
        assert "pepali arah" in text.lower()
        assert "tanggal jawa" in text.lower()
        assert "tanggal hijriyah" in text.lower()
        assert artifact.name == "1990-04-25-partner-2025-01-14.docx"
        del doc



def test_docx_export_can_embed_manual_trace_and_bibliography() -> None:
    with output_dir() as artifact_dir:
        artifact = write_profile_docx(
            target_date=date(1990, 4, 25),
            partner_date=date(2025, 1, 14),
            events=("nikah",),
            output_dir=artifact_dir,
            include_manual_calculation=True,
            include_bibliography=True,
        )

        doc = Document(artifact)
        text = _document_text(doc)
        assert "Detail perhitungan manual" in text
        assert "Detail Perhitungan Manual" in text
        assert "Offset hari dari epoch" in text
        assert "Pustaka dan basis aturan" in text
        assert "Pustaka dan Basis Aturan" in text
        assert "Kraton Yogyakarta" in text
        del doc



def test_docx_export_uses_unique_name_for_repeat_runs() -> None:
    with output_dir() as artifact_dir:
        first_artifact = write_profile_docx(
            target_date=date(1990, 4, 25),
            output_dir=artifact_dir,
        )
        second_artifact = write_profile_docx(
            target_date=date(1990, 4, 25),
            output_dir=artifact_dir,
        )

        assert first_artifact.name == "1990-04-25.docx"
        assert second_artifact.name == "1990-04-25-2.docx"



def test_docx_export_uses_unique_name_for_repeat_partner_runs() -> None:
    with output_dir() as artifact_dir:
        first_artifact = write_profile_docx(
            target_date=date(1990, 4, 25),
            partner_date=date(2025, 1, 14),
            output_dir=artifact_dir,
        )
        second_artifact = write_profile_docx(
            target_date=date(1990, 4, 25),
            partner_date=date(2025, 1, 14),
            output_dir=artifact_dir,
        )

        assert first_artifact.name == "1990-04-25-partner-2025-01-14.docx"
        assert second_artifact.name == "1990-04-25-partner-2025-01-14-2.docx"
