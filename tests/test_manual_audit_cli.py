from __future__ import annotations

from contextlib import contextmanager
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Iterator

from docx import Document

from research_system.commands.javanese_profile import main
from research_system.utils import javanese_profile_bibliography, manual_calculation_detail


@contextmanager
def cli_output_dir() -> Iterator[Path]:
    base_dir = Path("output")
    base_dir.mkdir(exist_ok=True)
    with TemporaryDirectory(dir=base_dir) as tmp_dir:
        yield Path(tmp_dir)


def _document_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    parts.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    return "\n".join(parts)


def test_manual_calculation_detail_tracks_core_steps() -> None:
    report = manual_calculation_detail("1990-04-25", partner="2025-01-14", events=("nikah",))

    assert report.weton_jawa == "Rebo Pon"
    assert any(step.title == "Offset hari dari epoch" for step in report.steps)
    assert any(step.title == "Neptu weton" and step.result == "14" for step in report.steps)
    assert any(
        step.title == "Tanggal Jawa dan Hijriyah"
        and "29 Pasa 1922 AJ" in step.result
        and "29 Ramadan 1410 H" in step.result
        for step in report.steps
    )
    assert any(step.title == "Jenjem gabungan" and step.result == "22 -> Netral" for step in report.steps)
    assert any(step.title == "Acara nikah" and step.section == "Hari baik" for step in report.steps)


def test_bibliography_keeps_external_and_internal_entries_visible() -> None:
    entries = javanese_profile_bibliography()
    external_only = javanese_profile_bibliography(include_internal=False)

    assert any(entry.topic == "Kalender Jawa Sultan Agungan" and entry.url for entry in entries)
    assert any(entry.topic == "Jenjem pasangan dan hari baik" and entry.source_kind == "internal" for entry in entries)
    assert external_only
    assert all(entry.source_kind == "external" for entry in external_only)


def test_main_writes_manual_trace_and_bibliography_to_docx(capsys) -> None:
    with cli_output_dir() as output_dir:
        main(
            [
                "--date",
                "1990-04-25",
                "--partner-date",
                "2025-01-14",
                "--events",
                "nikah",
                "--output-dir",
                str(output_dir),
                "--detail-perhitungan-manual",
                "--pustaka",
            ]
        )

        output = capsys.readouterr().out
        artifact = output_dir / "1990-04-25-partner-2025-01-14.docx"
        assert str(artifact) in output
        assert "Detail Perhitungan Manual" not in output
        assert "Pustaka dan Basis Aturan" not in output
        assert artifact.exists()

        doc = Document(artifact)
        text = _document_text(doc)
        assert "Detail Perhitungan Manual" in text
        assert "Offset hari dari epoch" in text
        assert "Pustaka dan Basis Aturan" in text
        assert "Kraton Yogyakarta" in text
